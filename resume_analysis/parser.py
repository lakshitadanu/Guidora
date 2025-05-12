import os
import re
from typing import Dict, List
import fitz  # PyMuPDF
from docx import Document
import magic
import pdfplumber
import docx

class ResumeParser:
    def __init__(self):
        self.skill_patterns = [
            r'(?i)(?:skills|technical skills|core competencies|expertise)[:|\n](.*?)(?=\n\n|\Z)',
            r'(?i)(?:proficient in|experienced with|familiar with)[:|\n](.*?)(?=\n\n|\Z)',
        ]
        
        self.education_patterns = [
            r'(?i)(?:education|academic background|qualifications)[:|\n](.*?)(?=\n\n|\Z)',
            r'(?i)(?:degree|bachelor|master|phd|diploma).*?(?=\n\n|\Z)',
        ]
        
        self.experience_patterns = [
            r'(?i)(?:experience|work history|professional experience)[:|\n](.*?)(?=\n\n|\Z)',
            r'(?i)(?:worked at|employed at|position at).*?(?=\n\n|\Z)',
        ]
        
        self.common_skills = {
            'Programming': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust'],
            'Web Development': ['html', 'css', 'react', 'angular', 'vue', 'django', 'flask', 'node.js', 'express'],
            'Database': ['sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'oracle'],
            'DevOps': ['docker', 'kubernetes', 'aws', 'azure', 'gcp', 'jenkins', 'git', 'ci/cd'],
            'Data Science': ['python', 'r', 'pandas', 'numpy', 'tensorflow', 'pytorch', 'scikit-learn'],
            'Cloud': ['aws', 'azure', 'gcp', 'cloud computing', 'serverless'],
            'Soft Skills': ['communication', 'leadership', 'teamwork', 'problem-solving', 'time management']
        }

    def parse_resume(self, file_path):
        """
        Parse a resume file (PDF or DOCX) and extract text and structured info.
        Args:
            file_path (str): Path to the resume file.
        Returns:
            dict: A dictionary containing extracted info (name, skills, education, experience, raw_text, file_type).
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        file_ext = os.path.splitext(file_path)[1].lower()
        if file_ext == '.pdf':
            text = self._extract_text_from_pdf(file_path)
        elif file_ext == '.docx':
            text = self._extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")

        # Clean and normalize the text
        cleaned_text = self._clean_text(text)

        # Extract structured info
        info = self._extract_info(cleaned_text)
        info['file_type'] = file_ext
        # Extract name (first non-empty line, likely at the top)
        lines = text.split('\n')
        name = ''
        for line in lines[:10]:  # Only look at the top 10 lines
            line = line.strip()
            # Skip section headers and lines that are all uppercase or too short
            if not line or line.isupper() or len(line.split()) < 2:
                continue
            # Allow letters, spaces, dots, and hyphens
            if re.match(r'^[A-Za-z .\-]+$', line):
                name = line
                break
        info['name'] = name
        return info

    def _extract_text_from_pdf(self, file_path):
        """
        Extract text from a PDF file using pdfplumber.
        Args:
            file_path (str): Path to the PDF file.
        Returns:
            str: Extracted text.
        """
        text = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
        return text

    def _extract_text_from_docx(self, file_path):
        """
        Extract text from a DOCX file using python-docx.
        Args:
            file_path (str): Path to the DOCX file.
        Returns:
            str: Extracted text.
        """
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text

    def _clean_text(self, text):
        """
        Clean and normalize the extracted text for LLM processing.
        Args:
            text (str): Raw text extracted from the resume.
        Returns:
            str: Cleaned and normalized text.
        """
        # Remove extra whitespace and newlines
        text = re.sub(r'\s+', ' ', text).strip()
        return text

    def _get_file_type(self, file_path: str) -> str:
        """Get MIME type of the file"""
        return magic.from_file(file_path, mime=True)

    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file using PyMuPDF"""
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        return text

    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        doc = Document(file_path)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])

    def _extract_sections(self, text):
        # Define common section headers
        section_headers = [
            "EDUCATION", "WORK EXPERIENCE", "EXPERIENCE", "PROJECTS", "CERTIFICATIONS", "SKILLS", "EXTRACURRICULAR", "LANGUAGE", "SOFT- SKILLS"
        ]
        # Build a regex pattern to split on section headers
        pattern = r'(' + '|'.join([re.escape(h) for h in section_headers]) + r')'
        # Split the text into sections
        splits = re.split(pattern, text, flags=re.IGNORECASE)
        # Map section name to content
        sections = {}
        current_section = None
        for part in splits:
            part = part.strip()
            if part.upper() in section_headers:
                current_section = part.upper()
                sections[current_section] = ""
            elif current_section:
                sections[current_section] += part + " "
        return sections

    def _extract_institutions(self, text, keywords):
        lines = text.split('\n')
        institutions = []
        for line in lines:
            for kw in keywords:
                if kw.lower() in line.lower():
                    # Try to extract just the institution/company name (first phrase)
                    name = line.split(',')[0].split('(')[0].split('-')[0].strip()
                    if name and name not in institutions:
                        institutions.append(name)
        return institutions

    def _extract_info(self, text: str) -> Dict:
        sections = self._extract_sections(text)
        # --- Education ---
        edu_section = sections.get("EDUCATION", "")
        education = []
        if edu_section:
            # Split by line, filter out empty lines
            education = [line.strip() for line in edu_section.split('\n') if line.strip()]
        else:
            # Fallback: look for lines with degree keywords
            degree_keywords = ["bachelor", "master", "phd", "b.tech", "m.tech", "b.e", "m.e", "mba", "b.sc", "m.sc", "school", "university", "college", "institute"]
            education = [line.strip() for line in text.split('\n') if any(kw in line.lower() for kw in degree_keywords)]
        # Only top 2
        education = education[:2]
        # --- Experience ---
        exp_section = sections.get("WORK EXPERIENCE", "") or sections.get("EXPERIENCE", "")
        experience = []
        if exp_section:
            experience = [line.strip() for line in exp_section.split('\n') if line.strip()]
        else:
            # Fallback: look for lines with job/role keywords
            job_keywords = ["engineer", "developer", "manager", "intern", "consultant", "analyst", "lead", "project", "trainee", "company", "technologies", "solutions", "systems", "labs"]
            experience = [line.strip() for line in text.split('\n') if any(kw in line.lower() for kw in job_keywords)]
        experience = experience[:2]
        # --- Skills ---
        extracted_skills = self._extract_skills(text)
        return {
            'skills': extracted_skills,
            'education': education,
            'experience': experience,
        }

    def _extract_skills(self, text: str) -> Dict[str, List[str]]:
        """Extract and categorize skills"""
        skills_found = {category: [] for category in self.common_skills.keys()}
        
        # Convert text to lowercase for case-insensitive matching
        text_lower = text.lower()
        
        # Check for each skill category
        for category, skills in self.common_skills.items():
            for skill in skills:
                if skill.lower() in text_lower:
                    skills_found[category].append(skill)
        
        # Remove empty categories
        return {k: v for k, v in skills_found.items() if v} 